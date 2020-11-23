import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
import xlsxwriter
import xlrd
import os


class DataModel(object):
    def __init__(self, e, i):
        self.entries = e
        self.items = i
        self.data = []
        self.activeEntry = ''
        self.income = 0.0
        self.outcome = 0.0
        self.pieData = {}

        self.incomeText = 'Prijem'
        self.savings = 0.0

    def plotPie(self, datain):
        # print(self.pieData)
        self.pieData = {key: val for key, val in datain.items() if val != 0.0}
        if self.incomeText in self.pieData:
            self.pieData[self.incomeText] = self.savings

        self.pieData['Uspora'] = self.pieData.pop('Prijem')
        _fig, ax = plt.subplots(figsize=(10, 4), subplot_kw=dict(aspect="equal"))
        data = list(self.pieData.values())
        keys = list(self.pieData.keys())
        wedges, _texts = ax.pie(data, wedgeprops=dict(width=0.5), startangle=-40)
        labels = [x.split()[-1] for x in keys]

        bbox_props = dict(boxstyle="square,pad=0.3", fc="w", ec="k", lw=0.72)
        kw = dict(arrowprops=dict(arrowstyle="-"), bbox=bbox_props, zorder=0, va="center")

        for i, p in enumerate(wedges):
            ang = (p.theta2 - p.theta1) / 2. + p.theta1
            y = np.sin(np.deg2rad(ang))
            x = np.cos(np.deg2rad(ang))
            horizontalalignment = {-1: "right", 1: "left"}[int(np.sign(x))]
            connectionstyle = "angle,angleA=0,angleB={}".format(ang)
            kw["arrowprops"].update({"connectionstyle": connectionstyle})
            ax.annotate(keys[i], xy=(x, y), xytext=(1.35 * np.sign(x), 1.4 * y),
                        horizontalalignment=horizontalalignment, **kw)

        ax.legend(wedges, labels, title='Type', loc='center left', bbox_to_anchor=(-1.0, 0, 0.5, 1))
        ax.set_title("Months summary", loc='left')
        return plt

    def setData(self, data):
        self.data = data

    def setActiveEntry(self, entry):
        self.activeEntry = entry

    def showGraph(self):
        print('showGraph')
        inputData = self.preprocessDataForPie()
        graph = self.plotPie(inputData)
        graph.show()

    def calcCZK(self, eur):
        return eur * 26.58

    def addTable(self, data, document):
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Value EUR'
        hdr_cells[2].text = 'Value CZK'
        hdr_cells[3].text = 'Description'
        for date, valueEU, valueCZ, desc in data:
            row_cells = table.add_row().cells
            row_cells[0].text = date
            row_cells[1].text = '{:.2f}'.format(valueEU) + ' EUR'
            row_cells[2].text = '{:.2f}'.format(valueCZ) + ' CZK'
            row_cells[3].text = desc

    def printSummary(self):
        file = str(self.activeEntry).split('_')
        file[2] = file[2].rstrip('.csv')
        plt.clf()
        memfile = BytesIO()
        inputData = self.preprocessDataForPie()
        figure = self.plotPie(inputData)
        figure.savefig(memfile)

        document = Document()
        document.add_heading(file[2].upper() + ' ' + file[1], 0)

        document.add_picture(memfile, width=Inches(7.75))

        document.add_paragraph('Income', style='Intense Quote')

        records = (
            ('5/2020', inputData['Prijem'], self.calcCZK(inputData['Prijem']), 'Celkovy prijem'),
        )
        self.addTable(records, document)
        print(inputData)

        document.add_paragraph('Expenses', style='Intense Quote')

        records = ()
        for key in inputData.keys():
            if self.incomeText not in key:
                if inputData[key] != 0.0:
                    record = ('5/2020', inputData[key], self.calcCZK(inputData[key]), key)
                    records = records + (record,)

        self.addTable(records, document)

        documentName = './reports/' + file[1] + '_' + file[2] + '.docx'
        document.save(documentName)
        memfile.close()

    def preprocessDataForPie(self):
        inputData = {}
        self.savings = 0.0
        for item in self.items:
            inputData[item] = 0.0
        for item in self.data:
            if 'Price EUR' not in item:
                try:
                    inputData[item[0]] += float(item[2])
                except KeyError:
                    print('key not found')
        for key in inputData.keys():
            if self.incomeText not in key:
                self.outcome += inputData[key]
        for item in self.data:
            if self.incomeText in item:
                print("naslo")
                self.savings += float(item[2])
        return inputData

    def showSummary(self):
        pass

    def readDataSummary(self):
        filePath = os.path.abspath(os.getcwd())
        os.system(filePath + '/reports/' + 'Summary.xlsx')

    def getDictionary(self, header):
        placeholder = []

        for item in header:  # TODO: find better way
            placeholder.append([])

        dictionary = dict(zip(header, placeholder))

        for col in self.data:
            if col[0] in dictionary.keys():
                dictionary[col[0]].append(float(col[2]))

        for item in dictionary:
            if not dictionary[item]:
                dictionary[item].append(float(0.0))

        max_length = 0
        for item in dictionary:
            size = len(dictionary[item])
            if size > max_length:
                max_length = size

        for item in dictionary:
            while len(dictionary[item]) < max_length:
                dictionary[item].append(float(0.0))

        return pd.DataFrame(data=dictionary)

    def saveDataSummary(self):
        headers = ['Nafta/Benzin', 'Byt', 'Jidlo', 'Automat', 'Hazard', 'Sport', 'Obleceni', 'Elektronika', 'Zahradka',
                   'Auto', 'Zabava', 'Lekarna', 'Prijem', 'Ostatni']

        df = self.getDictionary(header=headers)

        with pd.ExcelWriter("./reports/Expenses01.xlsx", engine="openpyxl", mode="a") as writer:
            df.to_excel(writer, startrow=0, startcol=0, sheet_name=self.activeEntry.rstrip('.csv').lstrip('_'))

        print('statistic saved')

    def webMonth(self):
        pass

    def setIncomeText(self, text):
        self.incomeText = text

    def getIncomeText(self):
        return self.incomeText

    def testCreator(self):
        print(self.entries)
        print(self.items)

    def testData(self):
        print(self.data)
